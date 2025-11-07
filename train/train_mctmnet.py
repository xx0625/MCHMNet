import os
import re
import numpy as np
import pandas as pd
import torch
import torch.nn as nn
import torch.optim as optim
import matplotlib.pyplot as plt
from sklearn.metrics import accuracy_score, f1_score, confusion_matrix
from tqdm import tqdm
import sys
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.styles import styles

# 确保中文显示正常
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]

sys.path.append(r'C:/xx/Chinese/main/models')
from MCTMNet import MCTMNet


def setup_seed(seed):
    np.random.seed(seed)
    torch.manual_seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)
    torch.backends.cudnn.deterministic = True
    torch.backends.cudnn.benchmark = False
    os.environ['PYTHONHASHSEED'] = str(seed)


setup_seed(42)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Using device: {device}")


def load_and_reshape(csv_path, session_id, label, n_timesteps=250, channel=61, window_type='hamming'):
    """
    加载数据并重塑为模型输入格式
    """
    df = pd.read_csv(csv_path)

    # 前61列是信息，之后是EEG数据，倒数第三列是身份
    data = df.iloc[:, :channel].values * 1e6
    identities = df.iloc[:, -3].values  # 身份信息（倒数第三列）

    # 创建带会话和标签信息的身份标识
    full_identities = [f"s{session_id}_id{int(identity)}_label{label}"
                       for identity in identities]

    # 选择窗函数
    if window_type == 'hamming':
        window = np.hamming(n_timesteps)
    elif window_type == 'hanning':
        window = np.hanning(n_timesteps)
    elif window_type == 'blackman':
        window = np.blackman(n_timesteps)
    else:
        window = np.ones(n_timesteps)

    reshaped_data = []
    reshaped_labels = []
    reshaped_identities = []

    # 滑动窗口分割数据
    for i in range(0, len(data) - n_timesteps + 1, n_timesteps):
        # 所有窗口都使用从文件名提取的同一标签
        windowed_segment = data[i:i + n_timesteps] * window[:, np.newaxis]
        reshaped_data.append(windowed_segment)
        reshaped_labels.append(label)  # 使用从文件名提取的标签
        reshaped_identities.append(full_identities[i])

    reshaped_data = np.array(reshaped_data)
    return reshaped_data.transpose(0, 2, 1), np.array(reshaped_labels), reshaped_identities


def prepare_data(train_sessions, test_session, base_path, n_timesteps=256, channel=14):
    """
    准备训练、验证和测试数据
    """
    train_X, train_y, train_identities = [], [], []
    val_X, val_y, val_identities = [], [], []

    # 加载训练会话数据
    for session in train_sessions:
        # 查找该会话的所有文件（格式如1_0.csv, 1_1.csv等）
        session_files = [f for f in os.listdir(base_path)
                         if f.startswith(f"{session}_") and f.endswith(".csv")]

        for file in session_files:
            try:
                # 分割文件名获取标签（假设格式为 "session_label.csv"）
                label = int(file.split('_')[1].split('.')[0])
                if label < 0 or label > 3:
                    print(f"警告: 文件 {file} 包含无效标签 {label}，已跳过")
                    continue
            except (IndexError, ValueError):
                print(f"警告: 无法从文件名 {file} 提取标签，已跳过")
                continue

            file_path = os.path.join(base_path, file)
            # 加载数据，传入从文件名提取的标签
            data, labels, identities = load_and_reshape(
                file_path, session, label, n_timesteps, channel
            )

            # 划分训练集和验证集（9:1）
            indices = np.random.permutation(len(data))
            num_val = int(len(data) * 0.1)

            val_X.append(data[indices[:num_val]])
            val_y.append(labels[indices[:num_val]])
            val_identities.extend([identities[i] for i in indices[:num_val]])

            train_X.append(data[indices[num_val:]])
            train_y.append(labels[indices[num_val:]])
            train_identities.extend([identities[i] for i in indices[num_val:]])

    # 合并训练数据
    train_X = np.vstack(train_X) if train_X else np.array([])
    train_y = np.hstack(train_y) if train_y else np.array([])
    train_identities = np.array(train_identities)

    # 合并验证数据
    val_X = np.vstack(val_X) if val_X else np.array([])
    val_y = np.hstack(val_y) if val_y else np.array([])
    val_identities = np.array(val_identities)

    # 加载测试会话数据
    test_X, test_y, test_identities = [], [], []
    test_files = [f for f in os.listdir(base_path)
                  if f.startswith(f"{test_session}_") and f.endswith(".csv")]

    for file in test_files:
        # 从文件名提取标签
        try:
            label = int(file.split('_')[1].split('.')[0])
            if label < 0 or label > 3:
                print(f"警告: 测试文件 {file} 包含无效标签 {label}，已跳过")
                continue
        except (IndexError, ValueError):
            print(f"警告: 无法从测试文件名 {file} 提取标签，已跳过")
            continue

        file_path = os.path.join(base_path, file)
        data, labels, identities = load_and_reshape(
            file_path, test_session, label, n_timesteps, channel
        )

        test_X.append(data)
        test_y.append(labels)
        test_identities.extend(identities)

    # 合并测试数据
    test_X = np.vstack(test_X) if test_X else np.array([])
    test_y = np.hstack(test_y) if test_y else np.array([])
    test_identities = np.array(test_identities)

    return train_X, train_y, train_identities, val_X, val_y, val_identities, test_X, test_y, test_identities


class EarlyStopping:
    def __init__(self, patience=7, verbose=False, delta=0, path='./best/checkpoint.pt'):
        self.patience = patience
        self.verbose = verbose
        self.counter = 0
        self.best_score = None
        self.early_stop = False
        self.val_loss_min = np.Inf
        self.delta = delta
        self.path = path

    def __call__(self, val_loss, model):
        score = -val_loss

        if self.best_score is None:
            self.best_score = score
            self.save_checkpoint(val_loss, model)
        elif score < self.best_score + self.delta:
            self.counter += 1
            print(f'EarlyStopping counter: {self.counter} out of {self.patience}')
            if self.counter >= self.patience:
                self.early_stop = True
        else:
            self.best_score = score
            self.save_checkpoint(val_loss, model)
            self.counter = 0

    def save_checkpoint(self, val_loss, model):
        if self.verbose:
            print(f'Validation loss decreased ({self.val_loss_min:.6f} --> {val_loss:.6f}).  Saving model ...')
        torch.save(model.state_dict(), self.path)
        self.val_loss_min = val_loss


def batch_predict(model, data, batch_size, device):
    """
    分批次预测函数，适用于验证集和测试集
    """
    model.eval()
    all_preds = []
    all_outputs = []
    with torch.no_grad():
        # 遍历所有批次
        for i in range(0, len(data), batch_size):
            batch_data = data[i:i + batch_size].to(device)
            outputs = model(batch_data)
            _, predicted = torch.max(outputs.data, 1)
            all_preds.extend(predicted.cpu().numpy())
            all_outputs.extend(outputs.cpu().numpy())
    return np.array(all_preds), np.array(all_outputs)


def save_results_to_word(model_name, params, session_results, avg_acc, avg_f1, std_acc, std_f1, save_path):
    """
    将实验结果保存到Word文档，格式与控制台输出一致
    """
    # 创建文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(10)

    # 添加标题
    title = doc.add_heading('EEGNet 模型实验结果', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加日期和时间
    doc.add_paragraph(f"实验日期: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    doc.add_paragraph(f"使用设备: {device}")
    doc.add_paragraph(f"模型版本: {model_name}")

    # 添加参数信息
    doc.add_heading('实验参数', level=1)
    param_text = "\n".join([f"{key}: {value}" for key, value in params.items()])
    doc.add_paragraph(param_text)
    doc.add_paragraph("")  # 空行

    # 添加各session结果，保持与控制台输出一致
    doc.add_heading('各Session测试结果', level=1)
    for session, result in session_results.items():
        # 测试session结果标题
        doc.add_paragraph(f"测试session {session} 结果 - 准确率: {result['acc']:.4f}, F1分数: {result['f1']:.4f}")

        # 混淆矩阵
        doc.add_paragraph("混淆矩阵:")
        cm_str = "\n".join([" ".join(map(str, row)) for row in result['cm']])
        para = doc.add_paragraph(cm_str)
        # 设置混淆矩阵缩进，使其更易读
        para.paragraph_format.left_indent = Cm(1)
        doc.add_paragraph("")  # 空行

    # 添加交叉验证最终结果，保持与控制台输出一致
    doc.add_heading('=== 留一session out 交叉验证最终结果 ===', level=1)
    for session in session_results.keys():
        doc.add_paragraph(
            f"Session {session} 作为测试集 - 准确率: {session_results[session]['acc']:.4f}, F1分数: {session_results[session]['f1']:.4f}")

    # 添加平均值和标准差
    doc.add_paragraph(f"平均准确率: {avg_acc:.4f} ± {std_acc:.4f}")
    doc.add_paragraph(f"平均F1分数: {avg_f1:.4f} ± {std_f1:.4f}")

    # 保存文档
    doc.save(save_path)
    print(f"\n实验结果已保存至: {save_path}")


def get_model_name():
    """从当前文件中提取EEGNet模型的导入文件名"""
    # 获取当前脚本的文件名
    current_script = os.path.basename(__file__)

    # 读取当前脚本内容
    with open(current_script, 'r', encoding='utf-8') as f:
        content = f.read()

    # 使用正则表达式查找导入语句
    pattern = r'from\s+(\w+)\s+import\s+EEGNet'
    match = re.search(pattern, content)

    if match:
        return match.group(1)  # 返回模型文件名
    else:
        return "unknown_model"  # 如果未找到，返回默认值


if __name__ == "__main__":
    base_path = "C:/xx/Chinese/data/preprocess_data"  # 修改为你的数据目录
    sessions = [1, 2, 3]  # 三个session
    n_timesteps = 250
    batch_size = 128
    max_epochs = 200
    patience = 20
    channel = 61
    all_acc, all_f1 = [], []
    session_results = {}  # 存储每个session的结果

    # 获取模型名称
    model_name = get_model_name()

    # 创建保存最佳模型和结果的目录
    os.makedirs('./best', exist_ok=True)
    os.makedirs('./results', exist_ok=True)

    for test_session in sessions:
        setup_seed(42)
        print(f"\n=== 留一session out - 测试session: {test_session} (Device: {device}) ===")
        train_sessions = [s for s in sessions if s != test_session]

        # 准备数据
        X_train, y_train, train_identities, X_val, y_val, val_identities, X_test, y_test, test_identities = prepare_data(
            train_sessions, test_session, base_path, n_timesteps, channel
        )
        print(X_train.shape)
        print(X_val.shape)
        print(X_test.shape)
        # 检查数据是否加载成功
        if len(X_train) == 0 or len(X_test) == 0:
            print(f"警告: 未能加载到session {test_session} 的有效数据，已跳过")
            continue

        # 转换为PyTorch张量
        X_train = torch.from_numpy(X_train).float().to(device)
        y_train = torch.from_numpy(y_train).long().to(device)
        X_val = torch.from_numpy(X_val).float().to(device)
        y_val = torch.from_numpy(y_val).long().to(device)
        X_test = torch.from_numpy(X_test).float().to(device)
        y_test = torch.from_numpy(y_test).long().to(device)

        # 初始化模型（四分类）
        model = MCTMNet(
            num_channels=channel,
            num_timesteps=n_timesteps,
            num_classes=4,
            noise_std=0.03,
            htfa_reduction=4,
            use_htfa=True,
            mamba_kernel_size=3
        ).to(device)
        # 定义优化器和损失函数
        optimizer = optim.Adam(model.parameters(), lr=0.001)
        criterion = nn.CrossEntropyLoss()  # 多分类交叉熵损失

        # 早停机制 - 保存路径包含模型名称和测试session
        checkpoint_path = f'./best/checkpoint_{model_name}_session{test_session}.pt'
        early_stopping = EarlyStopping(
            patience=patience,
            verbose=True,
            path=checkpoint_path
        )

        # 训练模型
        for epoch in tqdm(range(max_epochs), desc=f"训练进度 (测试session: {test_session})"):
            model.train()
            running_loss = 0.0
            total = 0
            correct = 0
            permutation = torch.randperm(X_train.size(0))

            for i in range(0, X_train.size(0), batch_size):
                indices = permutation[i:i + batch_size]
                batch_X, batch_y = X_train[indices], y_train[indices]
                optimizer.zero_grad()
                outputs = model(batch_X)
                loss = criterion(outputs, batch_y)
                loss.backward()
                optimizer.step()

                running_loss += loss.item()
                _, predicted = torch.max(outputs.data, 1)
                total += batch_y.size(0)
                correct += (predicted == batch_y).sum().item()

            epoch_loss = running_loss / (max(1, X_train.size(0) // batch_size))
            epoch_acc = correct / total if total > 0 else 0
            print(f'Epoch {epoch + 1}, 训练准确率: {epoch_acc:.4f}, 训练损失: {epoch_loss:.4f}')

            # 验证（分批次处理）
            model.eval()
            val_preds, val_outputs = batch_predict(model, X_val, batch_size, device)
            val_loss = criterion(torch.from_numpy(val_outputs).float().to(device), y_val)
            val_acc = accuracy_score(y_val.cpu().numpy(), val_preds)
            print(f'验证准确率: {val_acc:.4f}, 验证损失: {val_loss.item():.4f}')

            # 早停检查
            early_stopping(val_loss.item(), model)
            if early_stopping.early_stop:
                print(f"在第 {epoch + 1} 轮提前停止训练")
                break

        # 加载最佳模型并测试（分批次处理）
        model.load_state_dict(torch.load(checkpoint_path))
        test_preds, _ = batch_predict(model, X_test, batch_size, device)

        # 计算评估指标
        acc = accuracy_score(y_test.cpu().numpy(), test_preds)
        f1 = f1_score(y_test.cpu().numpy(), test_preds, average='macro')  # 多分类使用macro F1
        cm = confusion_matrix(y_test.cpu().numpy(), test_preds)

        all_acc.append(acc)
        all_f1.append(f1)
        # 保存当前session的结果
        session_results[test_session] = {
            'acc': acc,
            'f1': f1,
            'cm': cm
        }
        print(f"\n测试session {test_session} 结果 - 准确率: {acc:.4f}, F1分数: {f1:.4f}")
        print(f"混淆矩阵:\n{cm}")

    # 输出最终结果
    print("\n=== 留一session out 交叉验证最终结果 ===")
    for session, acc, f1 in zip(sessions, all_acc, all_f1):
        print(f"Session {session} 作为测试集 - 准确率: {acc:.4f}, F1分数: {f1:.4f}")

    if all_acc:  # 确保有有效数据
        avg_acc = np.mean(all_acc)
        avg_f1 = np.mean(all_f1)
        std_acc = np.std(all_acc)
        std_f1 = np.std(all_f1)

        print(f"平均准确率: {avg_acc:.4f} ± {std_acc:.4f}")
        print(f"平均F1分数: {avg_f1:.4f} ± {std_f1:.4f}")

        params = {
            'n_timesteps': n_timesteps,
            'batch_size': batch_size,
            'max_epochs': max_epochs,
            'patience': patience,
            'channel': channel,
            '设备': device
        }

        # 生成带时间戳的文件名，避免覆盖
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        result_path = f'./results/eegnet_results_{model_name}_{timestamp}.docx'
        save_results_to_word(model_name, params, session_results, avg_acc, avg_f1, std_acc, std_f1, result_path)
